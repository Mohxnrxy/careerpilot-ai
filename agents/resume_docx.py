from docx import Document


def create_resume_docx(content):

    doc = Document()

    doc.add_heading(
        "ATS Optimized Resume",
        level=1
    )

    doc.add_paragraph(content)

    filename = "improved_resume.docx"

    doc.save(filename)

    return filename